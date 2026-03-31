#!/usr/bin/env python3
"""
md_to_docx.py — Conversie Markdown → Word (.docx) cu formatare profesionala

Utilizare:
    python3 utils/md_to_docx.py utils/PREZENTARE_PROIECT_IDS-RS.md

Genereaza: utils/PREZENTARE_PROIECT_IDS-RS.docx
"""

import re
import sys
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================================
# Palete de culori
# =============================================================================
COLORS = {
    "primary":      RGBColor(0x1A, 0x3C, 0x6E),   # albastru inchis
    "secondary":    RGBColor(0x2B, 0x57, 0x9A),   # albastru mediu
    "accent":       RGBColor(0x3A, 0x6E, 0xA5),   # albastru deschis
    "accent_light": RGBColor(0xD6, 0xE4, 0xF0),   # albastru foarte deschis
    "text_dark":    RGBColor(0x2D, 0x2D, 0x2D),   # text principal
    "text_medium":  RGBColor(0x55, 0x55, 0x55),   # text secundar
    "text_light":   RGBColor(0x88, 0x88, 0x88),   # text subtil
    "code_red":     RGBColor(0xC0, 0x39, 0x2B),   # inline code
    "white":        RGBColor(0xFF, 0xFF, 0xFF),
    "bg_code":      "F2F4F7",                       # fundal code block
    "bg_quote":     "E8F0FE",                       # fundal blockquote
    "bg_header":    "1A3C6E",                       # fundal header tabel
    "separator":    RGBColor(0xBB, 0xBB, 0xBB),
}


# =============================================================================
# Helpers XML pentru formatare avansata
# =============================================================================

def set_cell_shading(cell, color_hex):
    """Seteaza culoarea de fundal a unei celule de tabel."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)
    cell._element.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, color_hex):
    """Seteaza culoarea de fundal a unui paragraf."""
    pPr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Seteaza border-urile unei celule."""
    tcPr = cell._element.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), val.get("sz", "4"))
            el.set(qn("w:color"), val.get("color", "CCCCCC"))
            el.set(qn("w:space"), "0")
            borders.append(el)
    tcPr.append(borders)


def add_page_number(paragraph):
    """Adauga numar de pagina in footer."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._element.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run3._element.append(fld_char_end)


def set_table_borders(table, color="CCCCCC", sz="4"):
    """Seteaza border-urile unui tabel complet."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
        borders.append(el)
    tblPr.append(borders)


# =============================================================================
# Stiluri document
# =============================================================================

def create_styles(doc):
    """Configureaza stilurile documentului."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = COLORS["text_dark"]
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.2

    for level, (size, color, sp_before, sp_after) in {
        1: (20, "primary",   Pt(28), Pt(12)),
        2: (15, "secondary", Pt(22), Pt(8)),
        3: (12, "accent",    Pt(16), Pt(6)),
        4: (11, "text_dark", Pt(12), Pt(4)),
    }.items():
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = COLORS[color]
        h.paragraph_format.space_before = sp_before
        h.paragraph_format.space_after = sp_after

    # Heading 2 cu linie sub el
    h2 = doc.styles["Heading 2"]
    h2_pPr = h2.element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "2B579A")
    bottom.set(qn("w:space"), "4")
    pBdr.append(bottom)
    h2_pPr.append(pBdr)

    # List Bullet
    lb = doc.styles["List Bullet"]
    lb.font.name = "Calibri"
    lb.font.size = Pt(11)
    lb.font.color.rgb = COLORS["text_dark"]
    lb.paragraph_format.space_after = Pt(3)
    lb.paragraph_format.line_spacing = 1.2


# =============================================================================
# Pagina de titlu
# =============================================================================

def add_title_page(doc, title, subtitle):
    """Adauga o pagina de titlu profesionala."""
    # Spatiu vertical
    for _ in range(6):
        doc.add_paragraph()

    # Linie decorativa sus
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 50)
    run.font.color.rgb = COLORS["secondary"]
    run.font.size = Pt(14)

    # Titlu principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]

    # Subtitlu
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(subtitle)
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = COLORS["text_medium"]

    # Linie decorativa jos
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 50)
    run.font.color.rgb = COLORS["secondary"]
    run.font.size = Pt(14)

    # Spatiu
    for _ in range(4):
        doc.add_paragraph()

    # Informatii document
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("DOCUMENT INTERN — PROPUNERE TEHNICA")
    run.font.size = Pt(10)
    run.font.color.rgb = COLORS["text_light"]
    run.font.small_caps = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datetime.now().strftime("%B %Y"))
    run.font.size = Pt(10)
    run.font.color.rgb = COLORS["text_light"]

    # Page break
    doc.add_page_break()


# =============================================================================
# Header si Footer
# =============================================================================

def setup_header_footer(doc, title_short):
    """Configureaza header si footer pe fiecare pagina."""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = hp.add_run(title_short)
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = COLORS["text_light"]
        run.font.small_caps = True

        run = hp.add_run("   |   Document intern")
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = COLORS["text_light"]

        # Linie sub header
        hp_pPr = hp._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "4")
        bot.set(qn("w:color"), "CCCCCC")
        bot.set(qn("w:space"), "6")
        pBdr.append(bot)
        hp_pPr.append(pBdr)

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = fp.add_run("— ")
        run.font.size = Pt(8)
        run.font.color.rgb = COLORS["text_light"]

        add_page_number(fp)

        run = fp.add_run(" —")
        run.font.size = Pt(8)
        run.font.color.rgb = COLORS["text_light"]


# =============================================================================
# Cuprins
# =============================================================================

def add_toc(doc):
    """Adauga un cuprins auto-actualizabil."""
    p = doc.add_heading("Cuprins", level=1)

    # Camp TOC (se actualizeaza in Word cu F9 sau click dreapta → Update Field)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    run2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instr)

    run3 = p.add_run()
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run3._element.append(fld_char_separate)

    run4 = p.add_run("(Click dreapta → Update Field pentru a genera cuprinsul)")
    run4.font.color.rgb = COLORS["text_light"]
    run4.font.size = Pt(9)
    run4.italic = True

    run5 = p.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run5._element.append(fld_char_end)

    doc.add_page_break()


# =============================================================================
# Formatare text inline
# =============================================================================

def add_formatted_text(paragraph, text):
    """Adauga text cu formatare inline (bold, italic, code, emoji)."""
    parts = re.split(r"(\*\*.*?\*\*|__.*?__|`[^`]+`|\*[^*]+\*|_[^_]+_)", text)

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("__") and part.endswith("__"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = COLORS["code_red"]
        elif (part.startswith("*") and part.endswith("*") and
              not part.startswith("**")):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif (part.startswith("_") and part.endswith("_") and
              not part.startswith("__")):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


# =============================================================================
# Elemente de document
# =============================================================================

def add_code_block(doc, lines):
    """Adauga un bloc de cod cu fundal gri si border."""
    code_text = "\n".join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)

    set_paragraph_shading(p, COLORS["bg_code"])

    # Border stanga albastru (accent)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:color"), "2B579A")
    left.set(qn("w:space"), "8")
    pBdr.append(left)
    pPr.append(pBdr)

    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = COLORS["text_dark"]


def add_table(doc, header_row, data_rows):
    """Adauga un tabel profesional cu header colorat."""
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Border-uri subtile
    set_table_borders(table, color="D0D0D0", sz="4")

    # Header — fundal albastru inchis, text alb
    for i, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, COLORS["bg_header"])
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        add_formatted_text(p, cell_text.strip())
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = COLORS["white"]
            run.font.name = "Calibri"

    # Data rows — alternanta alb / gri
    for row_idx, row_data in enumerate(data_rows):
        bg = "F8F9FA" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < cols:
                cell = table.rows[row_idx + 1].cells[col_idx]
                set_cell_shading(cell, bg)
                cell.text = ""
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                add_formatted_text(p, cell_text.strip())
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"

    # Spatiu dupa tabel
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def add_blockquote(doc, text):
    """Adauga un blockquote cu fundal albastru deschis si border stanga."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    set_paragraph_shading(p, COLORS["bg_quote"])

    # Border stanga
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:color"), "2B579A")
    left.set(qn("w:space"), "8")
    pBdr.append(left)
    pPr.append(pBdr)

    add_formatted_text(p, text)
    for run in p.runs:
        run.italic = True
        run.font.color.rgb = COLORS["text_medium"]
        run.font.size = Pt(10)


def add_separator(doc):
    """Adauga o linie subtire de separator."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 60)
    run.font.size = Pt(6)
    run.font.color.rgb = COLORS["separator"]


# =============================================================================
# Parsare tabel markdown
# =============================================================================

def parse_table_row(line):
    """Parseaza o linie de tabel markdown."""
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def is_separator_row(line):
    """Verifica daca linia este separator de tabel."""
    cleaned = line.strip().strip("|").replace("-", "").replace("|", "").replace(":", "").strip()
    return len(cleaned) == 0


# =============================================================================
# Conversia principala
# =============================================================================

def convert_md_to_docx(md_path, docx_path):
    """Converteste fisierul Markdown in Word cu formatare profesionala."""
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # Margini pagina
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    create_styles(doc)

    # Extragem titlul si subtitlul din primele linii
    title = "IDS-RS"
    subtitle = "Propunere de implementare"
    for line in lines[:5]:
        if line.startswith("# "):
            title = line[2:].strip()
            # Scoatem " — " si folosim ca subtitlu
            if " — " in title:
                parts = title.split(" — ", 1)
                title = parts[0].strip()
                subtitle = parts[1].strip()
            break

    # Pagina de titlu
    add_title_page(doc, title, subtitle)

    # Header si footer
    setup_header_footer(doc, title)

    # Cuprins
    add_toc(doc)

    # Procesam continutul
    i = 0
    first_h1_skipped = False

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Linie goala
        if not line.strip():
            i += 1
            continue

        # Separator ---
        if re.match(r"^---+\s*$", line.strip()):
            add_separator(doc)
            i += 1
            continue

        # Heading 1 — skipam primul (e titlul, deja pe cover page)
        if line.startswith("# ") and not line.startswith("## "):
            if not first_h1_skipped:
                first_h1_skipped = True
                i += 1
                continue
            text = line[2:].strip()
            p = doc.add_heading(level=1)
            add_formatted_text(p, text)
            i += 1
            continue

        # Headings
        if line.startswith("#### "):
            text = line[5:].strip()
            p = doc.add_heading(level=4)
            add_formatted_text(p, text)
            i += 1
            continue
        if line.startswith("### "):
            text = line[4:].strip()
            p = doc.add_heading(level=3)
            add_formatted_text(p, text)
            i += 1
            continue
        if line.startswith("## "):
            text = line[3:].strip()
            # Page break inainte de fiecare sectiune majora (H2)
            doc.add_page_break()
            p = doc.add_heading(level=2)
            add_formatted_text(p, text)
            i += 1
            continue

        # Code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            add_code_block(doc, code_lines)
            i += 1  # skip closing ```
            continue

        # Table
        if "|" in line and line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip("\n"))
                i += 1

            if len(table_lines) >= 2:
                header = parse_table_row(table_lines[0])
                data = []
                for tl in table_lines[1:]:
                    if not is_separator_row(tl):
                        data.append(parse_table_row(tl))
                add_table(doc, header, data)
            continue

        # Blockquote
        if line.strip().startswith("> "):
            quote_text = line.strip()[2:]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("> "):
                quote_text += " " + lines[i].strip()[2:]
                i += 1
            add_blockquote(doc, quote_text)
            continue

        # Bullet list
        if line.strip().startswith("- "):
            text = line.strip()[2:]
            i += 1
            while (i < len(lines) and lines[i].strip() and
                   not lines[i].strip().startswith("- ") and
                   not lines[i].strip().startswith("#") and
                   not lines[i].strip().startswith("|") and
                   not lines[i].strip().startswith("```") and
                   not lines[i].strip().startswith(">") and
                   not re.match(r"^---+\s*$", lines[i].strip()) and
                   (lines[i].startswith("  ") or lines[i].startswith("\t"))):
                text += " " + lines[i].strip()
                i += 1

            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, text)
            continue

        # Paragraph normal
        text = line.strip()
        i += 1
        while (i < len(lines) and lines[i].strip() and
               not lines[i].strip().startswith("#") and
               not lines[i].strip().startswith("- ") and
               not lines[i].strip().startswith("|") and
               not lines[i].strip().startswith("```") and
               not lines[i].strip().startswith("> ") and
               not re.match(r"^---+\s*$", lines[i].strip())):
            text += " " + lines[i].strip()
            i += 1

        p = doc.add_paragraph()
        add_formatted_text(p, text)

    doc.save(docx_path)
    print(f"DOCX generat: {docx_path}")


def main():
    if len(sys.argv) < 2:
        print("Utilizare: python3 md_to_docx.py <fisier.md>")
        sys.exit(1)

    md_path = sys.argv[1]
    if not os.path.exists(md_path):
        print(f"Fisierul nu exista: {md_path}")
        sys.exit(1)

    docx_path = os.path.splitext(md_path)[0] + ".docx"
    convert_md_to_docx(md_path, docx_path)


if __name__ == "__main__":
    main()
